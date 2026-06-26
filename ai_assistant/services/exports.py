from io import BytesIO

from django.core.files.base import ContentFile


def attach_draft_file(draft):
    if draft.output_format == "docx":
        content = _build_docx(draft.title, draft.content)
        if content:
            draft.file.save(f"draft_{draft.id}.docx", ContentFile(content), save=True)
    elif draft.output_format == "pdf":
        content = _build_pdf(draft.title, draft.content)
        if content:
            draft.file.save(f"draft_{draft.id}.pdf", ContentFile(content), save=True)
    return draft


def _build_docx(title, content):
    try:
        from docx import Document
    except Exception:
        return None

    buffer = BytesIO()
    document = Document()
    document.add_heading(title, level=1)
    for paragraph in content.split("\n"):
        document.add_paragraph(paragraph)
    document.save(buffer)
    return buffer.getvalue()


def _build_pdf(title, content):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except Exception:
        return None

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(50, y, title[:80])
    y -= 30
    pdf.setFont("Helvetica", 10)

    for paragraph in content.split("\n"):
        words = paragraph.split()
        line = ""
        for word in words:
            candidate = f"{line} {word}".strip()
            if len(candidate) > 95:
                pdf.drawString(50, y, line)
                y -= 14
                line = word
            else:
                line = candidate
            if y < 50:
                pdf.showPage()
                pdf.setFont("Helvetica", 10)
                y = height - 50
        if line:
            pdf.drawString(50, y, line)
            y -= 14
        y -= 8
        if y < 50:
            pdf.showPage()
            pdf.setFont("Helvetica", 10)
            y = height - 50

    pdf.save()
    return buffer.getvalue()
